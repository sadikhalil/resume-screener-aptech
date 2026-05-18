from docx import Document

def parse_docx(file_path: str) -> str:
    """Extract plain text from a DOCX file including tables."""
    doc = Document(file_path)
    parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts).strip()
